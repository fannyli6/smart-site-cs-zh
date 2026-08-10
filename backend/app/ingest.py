"""文档解析 + 切片 + 向量化入库。

支持格式：PDF / DOCX / XLSX / Markdown / TXT
切片策略：按段落聚合，目标 ~700 字/片，句子边界切分，100 字重叠，保留来源与章节线索。
"""
import io
import re
import uuid
from pathlib import Path

from . import config
from .store import add_chunks, delete_by_source

# ---------------------- 解析层 ----------------------
def _read_pdf(data: bytes) -> str:
    from pypdf import PdfReader
    try:
        reader = PdfReader(io.BytesIO(data))
        # 加密文档先尝试空密码解密，失败则明确报错
        if getattr(reader, "is_encrypted", False):
            try:
                reader.decrypt("")
            except Exception:
                raise ValueError("PDF 已加密，无法提取文本，请先解除密码后上传")
        text = "\n".join((p.extract_text() or "") for p in reader.pages)
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"PDF 解析失败：{e}")
    if not text.strip():
        raise ValueError(
            "PDF 未提取到文本，可能是扫描件/图片型 PDF（无文字层），请先做 OCR 或转成 Word/TXT 再上传"
        )
    return text


def _read_docx(data: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # 表格内容也纳入
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


def _read_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    out = []
    for ws in wb.worksheets:
        out.append(f"# 工作表：{ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                out.append(" | ".join(cells))
    return "\n".join(out)


def _read_text(data: bytes) -> str:
    return data.decode("utf-8", errors="ignore")


_READERS = {
    ".pdf": _read_pdf,
    ".docx": _read_docx,
    ".xlsx": _read_xlsx,
    ".xlsm": _read_xlsx,
    ".md": _read_text,
    ".txt": _read_text,
}


def parse_file(filename: str, data: bytes) -> str:
    ext = Path(filename).suffix.lower()
    reader = _READERS.get(ext)
    if not reader:
        # 兜底按文本处理
        return _read_text(data)
    return reader(data)


# ---------------------- 切片层 ----------------------
_CHUNK_SIZE = 700
_OVERLAP = 100
_HEADING_RE = re.compile(r"^(#{1,6}\s+.+|【.+】|第.+[章节条])")


def _split_sentences(text: str) -> list[str]:
    # 中文/英文句号、问号、换行都作为断句点
    parts = re.split(r"(?<=[。！？!?；;\n])", text)
    return [p for p in parts if p.strip()]


def chunk_text(text: str, source: str, domain: str, section_hint: str = "") -> list[dict]:
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if not text:
        return []
    sentences = _split_sentences(text)
    chunks = []
    buf = ""
    cur_section = section_hint
    for s in sentences:
        if _HEADING_RE.match(s.strip()):
            cur_section = s.strip()[:40]
        if len(buf) + len(s) > _CHUNK_SIZE and buf:
            chunks.append(_make_chunk(buf, source, domain, cur_section))
            # 重叠：取 buf 末尾若干字符作为下一片段开头
            buf = buf[-_OVERLAP:] + s
        else:
            buf += s
    if buf.strip():
        chunks.append(_make_chunk(buf, source, domain, cur_section))
    return chunks


def _make_chunk(text: str, source: str, domain: str, section: str) -> dict:
    return {
        "id": str(uuid.uuid4()),
        "text": text.strip(),
        "metadata": {
            "domain": domain,
            "source": source,
            "section": section,
        },
    }


# ---------------------- 入库层 ----------------------
def ingest_bytes(filename: str, data: bytes, domain: str) -> int:
    if domain not in config.VALID_DOMAINS:
        domain = config.DEFAULT_DOMAIN
    source = Path(filename).stem
    text = parse_file(filename, data)
    # 重新上传同名文件时，先清旧切片
    delete_by_source(source)
    chunks = chunk_text(text, source, domain)
    return add_chunks(chunks)


def ingest_markdown_string(title: str, text: str, domain: str) -> int:
    """直接灌入字符串知识（如种子 FAQ）。"""
    if domain not in config.VALID_DOMAINS:
        domain = config.DEFAULT_DOMAIN
    delete_by_source(title)
    chunks = chunk_text(text, title, domain)
    return add_chunks(chunks)
